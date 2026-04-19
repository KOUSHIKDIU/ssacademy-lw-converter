import streamlit as st
import openpyxl
import re
import io
from docx import Document

def parse_docx_final_polished(file):
    doc = Document(file)
    content = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
    
    tables_data = []
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in table_text:
                    table_text.append(text)
        tables_data.append(table_text)
        
    q_blocks = re.split(r'\n(?=Question No:\s*\d+)', content)
    q_blocks = [b for b in q_blocks if "Question No:" in b]
    
    data = []
    for idx, block in enumerate(q_blocks):
        item = {
            "Group": "Rheumatology", "Type": "TMC", "Question": "", "CorAns": "",
            "Answer1": "", "Answer2": "", "Answer3": "", "Answer4": "", "Answer5": "",
            "Answer6": "", "Answer7": "", "Answer8": "", "Answer9": "", "Answer10": "",
            "CorrectExplanation": "", "IncorrectExplanation": ""
        }
        
        lines = [line.strip() for line in block.split('\n') if line.strip()]
        
        idx_correct_ans, idx_explanation, idx_other = -1, -1, -1
        for i, line in enumerate(lines):
            if line.startswith("Correct Answer:"): idx_correct_ans = i
            elif line.startswith("Explanation:"): idx_explanation = i
            elif line.startswith("Other Options:"): idx_other = i
        
        if idx_correct_ans == -1: continue
            
        q_lines = lines[0 : idx_correct_ans - 5]
        if q_lines:
            q_no = q_lines[0].strip()
            if len(q_lines) >= 3:
                actual_q = q_lines[-1].strip()
                scenario = "<br>".join([l.strip() for l in q_lines[1:-1] if l.strip()])
                item["Question"] = f"<strong>{q_no}</strong><br>{scenario}<br><br><strong>{actual_q}</strong>".strip()
            elif len(q_lines) == 2:
                item["Question"] = f"<strong>{q_no}</strong><br><strong>{q_lines[1].strip()}</strong>".strip()
            else:
                item["Question"] = f"<strong>{q_no}</strong>".strip()
        
        opt_lines = lines[idx_correct_ans - 5 : idx_correct_ans]
        letters = ['A', 'B', 'C', 'D', 'E']
        for j in range(5):
            if j < len(opt_lines):
                item[f"Answer{j+1}"] = f"{letters[j]}. {opt_lines[j].strip()}".strip()
                
        ans_text = lines[idx_correct_ans]
        match = re.search(r'Correct Answer:\s*([A-E])', ans_text)
        if match:
            item["CorAns"] = ord(match.group(1).strip()) - ord('A') + 1
        
        idx_notes_marker = -1
        for i, line in enumerate(lines):
            if line.startswith("Notes:"): idx_notes_marker = i
            
        exp_end = idx_other if idx_other != -1 else (idx_notes_marker if idx_notes_marker != -1 else len(lines))
        exp_text = " ".join(lines[idx_explanation : exp_end]).replace("Explanation:", "").strip()
        
        other_end = idx_notes_marker if idx_notes_marker != -1 else len(lines)
        other_text = " ".join(lines[idx_other : other_end] if idx_other != -1 else []).replace("Other Options:", "").strip()
        
        other_parts = re.split(r'([A-E]\.)', other_text)
        others_formatted = ""
        if len(other_parts) > 1:
            for k in range(1, len(other_parts), 2):
                letter, desc = other_parts[k].strip(), other_parts[k+1].strip()
                col_idx = desc.find(':')
                if col_idx != -1:
                    others_formatted += f"<strong>{letter} {desc[:col_idx+1].strip()}</strong> {desc[col_idx+1:].strip()}<br>"
                else:
                    others_formatted += f"<strong>{letter}</strong> {desc}<br>"
        else:
            others_formatted = other_text.strip() + ("<br>" if other_text.strip() else "")
            
        notes_list = tables_data[idx] if idx < len(tables_data) else []
        notes_bullets = "<br>".join([f"* {n.strip()}" for n in notes_list if n.strip()])
        
        combined = f"<strong>Explanation:</strong><br>{exp_text}<br><br><strong>Other Options:</strong><br>{others_formatted}<strong>Notes:</strong><br>{notes_bullets}".strip()
        item["CorrectExplanation"], item["IncorrectExplanation"] = combined, combined
        data.append(item)
        
    return data

# --- Streamlit Web Interface ---
st.set_page_config(page_title="SsAcademy Converter", page_icon="📝")

# --- Hide Streamlit Branding ---
hide_streamlit_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            </style>
            """
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

st.title("LearnWorlds Bulk Importer")
st.write("Upload your nicely formatted Word Document, and download the Excel file ready for LearnWorlds!")

uploaded_file = st.file_uploader("Upload Word File (.docx)", type="docx")

if uploaded_file is not None:
    with st.spinner("Converting..."):
        try:
            # Parse the document
            questions_data = parse_docx_final_polished(uploaded_file)
            
            # Create Excel in memory
            wb = openpyxl.Workbook()
            ws_inst = wb.active
            ws_inst.title = "Instructions"
            ws_inst.append(["LearnWorlds Question Import Template Instructions"])
            ws_ex = wb.create_sheet(title="Examples")
            ws_ex.append(["Examples placeholder"])
            ws_q = wb.create_sheet(title="Questions")
            
            headers = ["Group", "Type", "Question", "CorAns", "Answer1", "Answer2", "Answer3", "Answer4", "Answer5", "Answer6", "Answer7", "Answer8", "Answer9", "Answer10", "CorrectExplanation", "IncorrectExplanation"]
            ws_q.append(headers)
            
            for row_idx, item in enumerate(questions_data, start=2):
                for col_idx, header in enumerate(headers, start=1):
                    val = item.get(header, "")
                    if isinstance(val, str): val = val.strip()
                    if val == "": val = None
                    ws_q.cell(row=row_idx, column=col_idx, value=val)
            
            # Save to an in-memory buffer
            output = io.BytesIO()
            wb.save(output)
            output.seek(0)
            
            st.success(f"Success! Converted {len(questions_data)} questions.")
            
            # Provide Download Button
            st.download_button(
                label="📥 Download Ready-to-Upload Excel",
                data=output,
                file_name="SsAcademy_LearnWorlds_Ready.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        except Exception as e:
            st.error(f"An error occurred: {e}")

# --- Developer Signature Footer ---
st.markdown("""
<div style="text-align: center; margin-top: 50px; padding-top: 20px; border-top: 1px solid #ddd;">
    <p style="margin-bottom: 5px; font-size: 14px; color: #666;">Design & Developed By</p>
    <p style="margin-bottom: 5px; font-size: 18px; font-weight: bold;">Koushik Sarkar</p>
    <p style="margin-bottom: 2px; font-size: 14px;">
        💬 WhatsApp: <a href="https://wa.me/8801732324131" target="_blank" style="text-decoration: none; color: #25D366; font-weight: bold;">+880 1732-324131</a>
    </p>
    <p style="margin-bottom: 2px; font-size: 14px;">
        📧 eMail: <a href="mailto:koushiksarkardiu@gmail.com" style="text-decoration: none; color: #0078D4; font-weight: bold;">koushiksarkardiu@gmail.com</a>
    </p>
</div>
""", unsafe_allow_html=True)
